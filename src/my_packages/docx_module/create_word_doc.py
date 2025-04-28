# pylint: disable=import-error, no-name-in-module, invalid-name,missing-module-docstring, missing-function-docstring, too-few-public-methods, too-many-branches, too-many-locals, too-many-statements
"""This module creates a Word document on the desktop with extracted values from a screenshot.
It uses the `read_image` function from the `ocr_module` to extract values such as ticket, mdn, and issue."""

import os

from docx import Document

from my_packages.ocr_module.read_image import read_image


def create_word_doc(title=None):
    """
    Creates a Word document on the desktop.
    The document filename will be the 'ticket' value extracted from a screenshot.
    Also adds ticket, mdn, and issue as paragraphs.

    Args:
        title (str): Title to add at the beginning of the document.

    Returns:
        str: Path to the created document.
    """
    # Call read_image and extract values
    result = read_image(0)
    ticket = result.get("ticket")
    mdn = result.get("mdn")
    issue = result.get("issue")

    if not ticket:
        raise ValueError("Ticket value not found in the image. Cannot create document.")

    # Create a new Document
    doc = Document()

    # Add title if provided
    if title:
        doc.add_heading(title, level=0)

    # Add paragraphs with only the extracted values
    doc.add_paragraph(ticket)
    doc.add_paragraph(mdn)
    doc.add_paragraph(issue)
    doc.add_paragraph("======================= ========================")  # separator

    # Get desktop path
    desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")

    # Create full path using ticket value as filename
    filename = f"{ticket}.docx"
    doc_path = os.path.join(desktop_path, filename)

    # Save the document
    doc.save(doc_path)

    return doc_path


# Example usage
if __name__ == "__main__":
    create_word_doc()
