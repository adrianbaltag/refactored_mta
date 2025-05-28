"""This script updates a Word document with images from a specified folder."""

import os

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches


def update_word_docx(indexes, width, clean_trailing=False):
    """_summary_: This function updates a Word document with images from a specified folder.

    Args:
        indexes (_type_): A list of indexes corresponding to the images to be added.
        width (_type_): The width of the images to be added in inches. --> constants saved in constant_variables.py
        clean_trailing (bool): Whether to remove trailing empty paragraphs before adding images.
                              Default is False to preserve existing content.
    """
    # Get the absolute path to the folder relative to the script's location
    desktop = os.path.join(os.path.expanduser("~"), "Desktop")
    base_dir = os.path.dirname(__file__)  # current dir of the script
    imgs_folder = os.path.join(base_dir, "..", "images")

    if not os.path.isdir(imgs_folder):
        print(f"'img' folder not found at: {imgs_folder}")
        return

    # Find the NRB document
    nrb_doc = None
    for file in os.listdir(desktop):
        if file.startswith("NRB") and file.endswith(".docx"):
            nrb_doc = os.path.join(desktop, file)
            break

    if not nrb_doc:
        print("No .docx file starting with 'NRB' found on Desktop.")
        return

    print(f"Using NRB file: {nrb_doc}")
    doc = Document(nrb_doc)

    # Get all image files in folder
    valid_exts = (".png", ".jpg", ".jpeg", ".PNG")
    # Get all image files in folder, filtering by valid extensions and sorting them alphabetically
    image_files = sorted([f for f in os.listdir(imgs_folder) if f.endswith(valid_exts)])

    # 🧹 Clean trailing empty paragraphs ONLY if explicitly requested
    # This prevents deleting existing content when adding new images
    if clean_trailing:
        print("🧹 Cleaning trailing empty paragraphs...")
        removed_count = 0
        while doc.paragraphs and not doc.paragraphs[-1].text.strip():
            p = doc.paragraphs[-1]._element
            p.getparent().remove(p)
            removed_count += 1
        if removed_count > 0:
            print(f"🧹 Removed {removed_count} empty paragraphs")

    for index in indexes:
        try:
            img_name = image_files[index]
            img_path = os.path.join(imgs_folder, img_name)

            paragraph = doc.add_paragraph()  # Creates a paragraph for the image
            run = paragraph.add_run()
            run.add_picture(img_path, width=Inches(width))  # Adjust size as needed
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # or CENTER

            print(f"✅ Added: {img_path}")
        except IndexError:
            print(f"❌ No image at index {index} in {imgs_folder}")

    doc.save(nrb_doc)
    print("✅ Document updated successfully.")


if __name__ == "__main__":
    # Example usage scenarios:

    # Scenario 1: First time adding images to a fresh document
    # Use clean_trailing=True to remove any extra empty paragraphs
    # update_word_docx([1], 2, clean_trailing=True)

    # Scenario 2: Adding more images to an existing document with content
    # Use clean_trailing=False (default) to preserve existing content
    # update_word_docx([3, 4], width=5)
    update_word_docx([4, 5], width=6)

    # Scenario 3: Multiple calls - clean only on first call
    # update_word_docx([1, 2], width=3, clean_trailing=True)  # First run - clean
    # update_word_docx([3, 4], width=5)  # Second run - don't clean
