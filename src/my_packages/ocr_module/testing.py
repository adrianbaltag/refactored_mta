def update_word_text(content=None, doc_name_prefix="NRB"):
    """Updates a Word document with content dynamically while preserving existing content.

    Args:
        content (any, optional): Content to add to the document.
            Can be string, list, dict, or nested structure of these types.
        doc_name_prefix (str, optional): Prefix of the document filename to search for.
            Defaults to "NRB".

    Returns:
        bool: True if update was successful, False otherwise.
    """
    import os

    from docx import Document

    try:
        # Find the document on the desktop
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")

        # Find the document
        doc_path = None
        for file in os.listdir(desktop):
            if file.startswith(doc_name_prefix) and file.endswith(".docx"):
                doc_path = os.path.join(desktop, file)
                break

        if not doc_path:
            print(f"No .docx file starting with '{doc_name_prefix}' found on Desktop.")
            return False

        print(f"Using file: {doc_path}")
        doc = Document(doc_path)

        # Helper function to add content recursively
        def add_content(content, level=0):
            if content is None:
                return

            # String - add as paragraph
            if isinstance(content, str):
                para = doc.add_paragraph()
                para.add_run(content)
                return

            # List or tuple - add as bullet points or sections
            elif isinstance(content, (list, tuple)):
                # Check if it's a heading-content pair
                if (
                    len(content) == 2
                    and isinstance(content[0], str)
                    and not isinstance(content[1], (list, tuple, dict))
                ):
                    # This looks like a (heading, content) pair
                    heading, text = content
                    doc.add_heading(heading, level=level + 1)
                    para = doc.add_paragraph()
                    para.add_run(str(text))
                else:
                    # Regular list - add each item
                    for item in content:
                        if isinstance(item, (list, tuple, dict)):
                            # Complex item - recurse
                            add_content(item, level + 1)
                        else:
                            # Simple item - add as bullet point
                            bullet = doc.add_paragraph(style="List Bullet")
                            bullet.add_run(str(item))
                return

            # Dictionary - add as table
            elif isinstance(content, dict):
                if content:  # Only create table if dict has items
                    table = doc.add_table(rows=len(content) + 1, cols=2)
                    table.style = "Table Grid"

                    # Add headers
                    header_cells = table.rows[0].cells
                    header_cells[0].text = "Item"
                    header_cells[1].text = "Value"

                    # Make header bold
                    for cell in header_cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.bold = True

                    # Add data rows
                    for i, (key, value) in enumerate(content.items()):
                        row = table.rows[i + 1].cells
                        row[0].text = str(key)

                        # Check if the value is complex
                        if isinstance(value, (list, tuple, dict)):
                            # For complex values, add a placeholder and then add content after table
                            row[1].text = "[See details below]"

                            # Add the complex content after the table
                            doc.add_paragraph()  # Add spacing
                            doc.add_heading(f"Details for {key}:", level=level + 2)
                            add_content(value, level + 2)
                        else:
                            row[1].text = str(value)

                    # Add space after table
                    doc.add_paragraph()
                return

            # Any other type - convert to string
            else:
                para = doc.add_paragraph()
                para.add_run(str(content))

        # Add a blank line before new content if the document isn't empty
        if len(doc.paragraphs) > 0 and doc.paragraphs[-1].text.strip():
            doc.add_paragraph()

        # Add the content to document
        add_content(content)

        # Save the document
        doc.save(doc_path)
        print(f"✅ Document '{doc_path}' updated successfully.")
        return True

    except Exception as e:
        print(f"❌ Error updating document: {str(e)}")
        return False


if __name__ == "__main__":
    # Example usage
    content = {
        "Introduction": "This is an introduction.",
        "Data": [1, 2, 3],
        "Details": {
            "Item 1": "Value 1",
            "Item 2": [4, 5, 6],
            "Item 3": {"Subitem A": "Value A", "Subitem B": "Value B"},
        },
    }
    update_word_text(content)
